import os
import re
import requests
import hashlib
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Path to the markdown file
MD_FILE = r'D:\tuphap\khodldc\dldc_1\tailieu\docs\compomennt.md'
OUTPUT_DOCX = r'D:\tuphap\khodldc\dldc_1\tailieu\docs\compomennt_v2.docx'
TEMP_DIR = r'D:\tuphap\khodldc\dldc_1\tailieu\docs\temp_images'

if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

def download_image(url):
    try:
        # Handle Iconify SVGs - convert to PNG
        if 'iconify.design' in url:
            # Clean up color parameter for Iconify PNG API
            url = url.replace('%23', '').replace('#', '')
            if '.svg' in url:
                url = url.replace('.svg', '.png')
            
            # Ensure / instead of : for prefix:name
            parts = url.split('.design/')
            if len(parts) > 1 and ':' in parts[1]:
                icon_path = parts[1].replace(':', '/', 1)
                url = parts[0] + '.design/' + icon_path
        
        url_hash = hashlib.md5(url.encode()).hexdigest()
        filename = f"img_{url_hash}.png"
        filepath = os.path.join(TEMP_DIR, filename)
        
        if not os.path.exists(filepath):
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code == 200:
                with open(filepath, 'wb') as f:
                    f.write(response.content)
            else:
                return None
        return filepath
    except Exception:
        return None

def set_cell_background(cell, fill):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_content_to_paragraph(paragraph, text, img_width=Inches(0.5)):
    # Clear paragraph
    paragraph.text = ""
    
    # Robust regex
    html_img_pattern = r'<img\s+[^>]*?src\s*=\s*["\'](?P<html_url>[^"\']+)["\'][^>]*?>'
    md_img_pattern = r'!\[(?P<md_alt>.*?)\]\((?P<md_url>.*?)\)'
    combined_pattern = re.compile(f'({html_img_pattern})|({md_img_pattern})', re.IGNORECASE | re.DOTALL)
    
    last_idx = 0
    matches = list(combined_pattern.finditer(text))
    
    if not matches:
        paragraph.add_run(text)
        return

    for match in matches:
        before_text = text[last_idx:match.start()]
        if before_text:
            paragraph.add_run(before_text)
        
        img_url = match.group('html_url') or match.group('md_url')
        if img_url:
            img_path = download_image(img_url)
            if img_path:
                try:
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=img_width)
                except:
                    paragraph.add_run(f" [Image Error] ")
            else:
                paragraph.add_run(f" [Link: {img_url}] ")
        
        last_idx = match.end()
    
    remaining_text = text[last_idx:]
    if remaining_text:
        paragraph.add_run(remaining_text)

def convert():
    if not os.path.exists(MD_FILE):
        print(f"Error: {MD_FILE} not found.")
        return

    doc = Document()
    
    try:
        style = doc.styles['Normal']
        style.font.name = 'Inter'
        style.font.size = Pt(11)
    except:
        pass

    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
            i += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            i += 1
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            i += 1
        elif line.startswith('|'):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|\s*:?-+:?\s*\|', row_line):
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    table_data.append(cells)
                i += 1
            
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for r in range(rows):
                    for c in range(cols):
                        cell = table.rows[r].cells[c]
                        content = table_data[r][c]
                        paragraph = cell.paragraphs[0]
                        add_content_to_paragraph(paragraph, content, img_width=Inches(0.25) if r > 0 else Inches(0.5))
                        if r == 0:
                            set_cell_background(cell, "F1F5F9")
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.bold = True
        elif line.startswith('---'):
            doc.add_paragraph("_" * 50)
            i += 1
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_content_to_paragraph(p, line[2:])
            i += 1
        elif line.startswith('<div'):
            while i < len(lines) and not lines[i].strip().endswith('</div>'):
                i += 1
            i += 1
            doc.add_paragraph("[Ví dụ giao diện]", style='Caption')
        else:
            p = doc.add_paragraph()
            add_content_to_paragraph(p, line)
            i += 1

    doc.save(OUTPUT_DOCX)
    print(f"Exported to: {OUTPUT_DOCX}")

if __name__ == "__main__":
    convert()
